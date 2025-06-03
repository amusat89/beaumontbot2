# -*- coding: utf-8 -*-
# LabMate - ISO 15189:2022 Compliant Laboratory Assistant

from docx import Document
import streamlit as st
from pathlib import Path
import json
import time
import pandas as pd
from chatlas import ChatGoogle
import hashlib
import os
import sys
from datetime import datetime

# -------------------------------
# Global Configuration
# -------------------------------
BASE_DIR = Path(__file__).parent.absolute()  # Changed to absolute() for cloud compatibility
LAB_DOCS_DIR = BASE_DIR / "lab_docs"
DEPARTMENT_FILES = {
    "In House Test": "In_House.docx",
    "External Referral": "ExternallyReferrredTests.docx"
}

# -------------------------------
# Debugging Functions
# -------------------------------
def verify_environment():
    """Debug environment and document verification"""
    debug_info = {
        "Python Version": sys.version,
        "Current Directory": str(Path.cwd()),
        "Base Directory": str(BASE_DIR),
        "LAB_DOCS_DIR": {
            "Path": str(LAB_DOCS_DIR),
            "Exists": LAB_DOCS_DIR.exists(),
            "Files Found": [f.name for f in LAB_DOCS_DIR.glob("*")] if LAB_DOCS_DIR.exists() else []
        },
        "Required Files": list(DEPARTMENT_FILES.values()),
        "Missing Files": [
            f for f in DEPARTMENT_FILES.values() 
            if not (LAB_DOCS_DIR / f).exists()
        ]
    }
    
    with st.expander("Environment Debug Information"):
        st.json(debug_info)
    
# -------------------------------
# Modified Debugging Functions
# -------------------------------
def verify_environment():
    """Silent environment verification"""
    if not LAB_DOCS_DIR.exists():
        st.error(f"Critical Error: Lab docs directory not found at {LAB_DOCS_DIR}")
        st.stop()
    
    missing_files = [f for f in DEPARTMENT_FILES.values() if not (LAB_DOCS_DIR / f).exists()]
    if missing_files:
        st.error(f"Missing required documents: {', '.join(missing_files)}")
        st.stop()

# -------------------------------
# Configuration
# -------------------------------
st.set_page_config(
    page_title="LabMate Pro",
    page_icon="🧬",
    layout="centered",
    initial_sidebar_state="expanded"
)

# -------------------------------
# Caching and Performance
# -------------------------------
@st.cache_resource(ttl=3600, show_spinner=False)
def load_all_departments():
    """Pre-load and cache all department documents with text+tables"""
    return {
        dept: process_docx(LAB_DOCS_DIR / filename)
        for dept, filename in DEPARTMENT_FILES.items()
    }

# -------------------------------
# Enhanced Data Processing
# -------------------------------
def process_docx(file_path: Path):
    """Process both text content and tables with context preservation"""
    try:
        doc = Document(file_path)
        
        text_content = []
        current_section = {"title": "Introduction", "content": []}
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                if para.style.name.startswith('Heading') or (para.runs and para.runs[0].bold):
                    if current_section["content"]:
                        text_content.append(current_section)
                    current_section = {
                        "title": text,
                        "content": [],
                        "subsections": []
                    }
                else:
                    current_section["content"].append(text)
        
        if current_section["content"]:
            text_content.append(current_section)
        
        return {
            "sections": text_content,
            "tables": [process_table(table) for table in doc.tables]
        }
    except Exception as e:
        st.error(f"DOCUMENT PROCESSING ERROR: {str(e)}")
        st.stop()

def process_table(table):
    """Convert DOCX tables to structured format with context"""
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    return {
        "context": get_table_context(table),
        "headers": headers,
        "rows": [
            {headers[i]: cell.text.strip() 
             for i, cell in enumerate(row.cells)}
            for row in table.rows[1:]
        ]
    }

def get_table_context(table):
    """Extract explanatory text preceding tables"""
    context = []
    prev_element = table._element.getprevious()
    while prev_element is not None and prev_element.tag.endswith('p'):
        text = prev_element.text.strip()
        if text:
            context.append(text)
        prev_element = prev_element.getprevious()
    return "\n".join(reversed(context)) if context else ""

# -------------------------------
# Prompt Engineering (Corrected)
# -------------------------------
def generate_system_prompt(selected_dept, dept_content):
    """Generate ISO-compliant prompt with proper escaping"""
    try:
        sections_md = []
        for section in dept_content['sections']:
            section_content = chr(10).join(section['content'])
            sections_md.append(f"### {section['title']}{chr(10)}{section_content}")
        
        sections_md = chr(10).join(sections_md)
        
        tables_md = []
        for i, table in enumerate(dept_content['tables']):
            tables_md.append(f"**Table {i+1}:**{chr(10)}{table_to_markdown(table)}")
        tables_md = chr(10).join(tables_md)
        
        return f"""
# Beaumont Hospital {selected_dept} Assistant Persona
**Role**: Senior Laboratory Specialist with 15 years experience
**Communication Style**: 
- Professionally friendly
- Clarifies ambiguities with numbered options
- Anticipates follow-up questions
- Uses natural medical terminology

## Response Protocol (ISO 15189 §7.11)

### Structured Response Format
- **Test Name**: {{test_name}}
- **Mnemonics**: {{mnemonics}}
- **Specimen**: {{sample_type}} ({{label_code}})
- **Container**: {{container_color}}
- **Blood Type**: {{SEP/WB}}
- **Storage**: {{storage}}
- **Minimum Volume**: {{volume}}
- **Method**: {{methodology}}
- **Reference Range**: {{reference_values}}
- **Turnaround**: {{turnaround_time}}
- **Notes**: {{critical_info}}
- **Department**: {{laboratory_department}}

*Only include fields present in the relevant table.*

## Department Reference Tables
{tables_md}

### Protocol Sections 
{sections_md} 

### Interactive Guidelines

1. For ambiguous specimen requests:
   "Which specimen type? Please select:
   1. Blood (EDTA)
   2. Urine (Sterile container)
   3. CSF (Tube 1)"

2. For test subtypes:
   "Which variant? Choose:
   1. Fasting Glucose
   2. Random Glucose
   3. 2hr Postprandial"

3. For demographic-dependent ranges:
   "For accurate ranges, please provide:
   1. Patient age
   2. Biological sex
   3. Pregnancy status (if applicable)"

4. For General information:
    "Would you like to know about:
    1. Location of Departments
    2. Contact information
    3. Department Opening Hours"
    

### Decision Logic
- Always present options as numbered lists
- Map user's number selection to exact protocol terms
- Confirm selection before proceeding: 
  "You selected: 1 (Blood). Processing blood test parameters..."

### Prohibited Formats
❌ No JSON/XML 
❌ No unnumbered options
❌ No assumptions without confirmation

### Example Workflow
User: "I need glucose test info"
Assistant: "Which specimen type? 
1. Plasma (Lithium Heparin)
2. Whole Blood 
3. CSF"
[User selects 1]
"Processing Plasma Glucose test... [structured response]"
"""
    except Exception as e:
        st.error(f"PROMPT GENERATION ERROR: {str(e)}")
        st.stop()

def table_to_markdown(table_data):
    """Convert table data to markdown format"""
    headers = table_data['headers']
    rows = table_data['rows']
    
    md = f"| {' | '.join(headers)} |\n"
    md += f"| {' | '.join(['---']*len(headers))} |\n"
    
    for row in rows:
        values = [str(row.get(h, '')) for h in headers]
        md += f"| {' | '.join(values)} |\n"
    
    return md

# -------------------------------
# Query Processing & Audit
# -------------------------------
def handle_query_type(prompt):
    """Classify queries for processing"""
    prompt = prompt.lower()
    test_keywords = ["test", "assay", "panel", "profile", "reference"]
    proc_keywords = ["procedure", "step", "guide", "handle", "process"]
    interp_keywords = ["interpret", "result", "significance", "means"]
    
    if any(k in prompt for k in test_keywords):
        return "TEST_PARAMETERS"
    elif any(k in prompt for k in proc_keywords):
        return "PROCEDURE"
    elif any(k in prompt for k in interp_keywords):
        return "INTERPRETATION"
    return "GENERAL"

def log_audit_entry(prompt, response, department):
    """ISO-compliant audit logging"""
    entry = {
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "session": st.session_state.session_id,
        "department": department,
        "query": prompt,
        "response": response,
        "hash": hashlib.sha256(response.encode()).hexdigest(),
        "compliance_check": check_compliance(response)
    }
    
    if "audit_log" not in st.session_state:
        st.session_state.audit_log = []
    st.session_state.audit_log.append(entry)
    
    with open("audit_log.csv", "a") as f:
        f.write(f"{json.dumps(entry)}\n")

def check_compliance(response):
    """Validate ISO 15189 compliance"""
    required_keys = ["Test Name", "Specimen", "Method"]
    return any(key in response for key in required_keys)

# -------------------------------
# UI Components
# -------------------------------
def render_department_selector(departments):
    """Department selector with lazy loading"""
    with st.sidebar:
        selected = st.selectbox("Select Department", list(departments.keys()))
        st.markdown("---")
        st.caption(f"ISO 15189:2022 Validated Protocols - {selected}")
    return selected

def display_chat_history():
    """Virtualized chat history for performance"""
    container = st.container()
    with container:
        for msg in st.session_state.messages[-10:]:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

# -------------------------------
# Main Application
# -------------------------------
def main():
    verify_environment()  # First check environment
    
    # Initialize session
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    if "session_id" not in st.session_state:
        st.session_state.session_id = hashlib.sha256(
            f"{datetime.now().isoformat()}_{os.getpid()}".encode()
        ).hexdigest()[:12]
    
    st.title("🧬 LabMate Pro")

    URL = "https://www.beaumont.ie/themes/custom/beaumont_barrio/logo.png"
    st.logo("https://www.beaumont.ie/themes/custom/beaumont_barrio/logo.png", link=URL, 
    size="large")
    st.markdown("**ISO 15189:2022 & LP-GEN-0016 Compliant Laboratory Assistant**")

    try:
        departments = load_all_departments()
    except Exception as e:
        st.error(f"Failed to load departments: {str(e)}")
        st.stop()
    
    selected_dept = render_department_selector(departments)
    
    if "current_dept" not in st.session_state or st.session_state.current_dept != selected_dept:
        st.session_state.current_dept = selected_dept
        st.session_state.messages = []
        st.session_state.chat = ChatGoogle(
            system_prompt=generate_system_prompt(selected_dept, departments[selected_dept]),
            api_key="AIzaSyAmO2k5pPLMsKKvuY79vlOrpYKQ-aC5Y74"
        )

    display_chat_history()
    
    if prompt := st.chat_input("Hi, How can i help you..."):
        handle_user_input(prompt, selected_dept, departments)

def handle_user_input(prompt, department, departments):
    """Process and route user queries"""
    query_type = handle_query_type(prompt)
    
    with st.chat_message("user"):
        st.markdown(prompt)
    
    st.session_state.messages.append({"role": "user", "content": prompt})
    
    with st.chat_message("assistant"):
        try:
            # Combine context with the query
            full_query = f"""Department: {department}
            Query Type: {query_type}
            User Question: {prompt}
            
            Answer using these resources:
            - Sections: {json.dumps(departments[department]['sections'][:3])} 
            - Tables: {json.dumps(departments[department]['tables'][:1])}
            """
            
            response_stream = st.session_state.chat.stream(full_query)
            
            response_placeholder = st.empty()
            full_response = []
            
            for chunk in response_stream:
                processed_chunk = format_response(chunk)
                full_response.append(processed_chunk)
                response_placeholder.markdown("".join(full_response) + "▌")
                time.sleep(0.02)
            
            final_response = "".join(full_response)
            response_placeholder.markdown(final_response)
            st.session_state.messages.append({"role": "assistant", "content": final_response})
            
            log_audit_entry(prompt, final_response, department)
            
        except Exception as e:
            handle_error(e)

def format_response(response):
    """Convert any JSON structures to proper formatting"""
    if "{" in response and "}" in response:
        try:
            json_str = response[response.find("{"):response.rfind("}")+1]
            data = json.loads(json_str)
            return "\n".join([f"**{k}:** {v}" for k, v in data.items()])
        except:
            return response
    return response

def handle_error(e):
    """Standardized error handling"""
    error_msg = f"SYSTEM ERROR: {str(e)}"
    st.error(error_msg)
    st.session_state.messages.append({
        "role": "assistant", 
        "content": "Error processing request - contact Lab Support"
    })
    log_audit_entry("SYSTEM ERROR", error_msg, "GLOBAL")

if __name__ == "__main__":
    main()
